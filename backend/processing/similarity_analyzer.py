"""
Similarity Analysis System for Word Documents
Implements TF-IDF + Semantic similarity checking for exam papers
"""

import os
import json
import time
import shutil
import numpy as np
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Any
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import re
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentSimilarityAnalyzer:
    def __init__(self, temp_base_dir: str = "similarityCheck/similarity_temp"):
        self.temp_base_dir = temp_base_dir
        self.similarity_threshold = 0.7
        self.tfidf_vectorizer = TfidfVectorizer(
            stop_words='english',
            max_features=5000,
            ngram_range=(1, 2),
            lowercase=True,
            strip_accents='unicode'
        )
        
        # Ensure temp directory exists
        os.makedirs(temp_base_dir, exist_ok=True)
    
    def create_session(self) -> str:
        """Create a new similarity analysis session"""
        timestamp = int(time.time())
        random_suffix = os.urandom(4).hex()
        session_id = f"similarity_{timestamp}_{random_suffix}"
        
        session_path = os.path.join(self.temp_base_dir, session_id)
        os.makedirs(session_path, exist_ok=True)
        os.makedirs(os.path.join(session_path, "uploaded_files"), exist_ok=True)
        os.makedirs(os.path.join(session_path, "extracted_content"), exist_ok=True)
        os.makedirs(os.path.join(session_path, "analysis_results"), exist_ok=True)
        
        # Create session info
        session_info = {
            "session_id": session_id,
            "created_at": datetime.now().isoformat(),
            "files_uploaded": [],
            "analysis_completed": False,
            "expires_at": (datetime.now() + timedelta(hours=2)).isoformat(),
            "session_metadata": {
                "temp_base_dir": self.temp_base_dir,
                "similarity_threshold": self.similarity_threshold,
                "tfidf_max_features": 5000,
                "session_structure": {
                    "uploaded_files": "uploaded_files/",
                    "extracted_content": "extracted_content/",
                    "analysis_results": "analysis_results/"
                }
            }
        }
        
        try:
            with open(os.path.join(session_path, "session_info.json"), 'w', encoding='utf-8') as f:
                json.dump(session_info, f, indent=2, ensure_ascii=False)
            logger.info(f"Created similarity session: {session_id}")
        except Exception as e:
            logger.error(f"Failed to create session info file: {str(e)}")
            # Clean up the directory if session creation fails
            if os.path.exists(session_path):
                shutil.rmtree(session_path)
            raise
        
        return session_id
    
    def extract_questions_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract questions and answers from Word document"""
        try:
            doc = Document(file_path)
            questions_by_type = {
                'mcq': [],
                'true_false': [],
                'matching': [],
                'short_answer': [],
                'long_answer': []
            }
            questions = []
            all_paragraphs = []
            current_section = None
            current_question = ""
            # Common exam instructions to ignore
            instruction_patterns = [
                # Basic instructions
                r'answer.*questions.*using.*pencil',
                r'write.*id.*name.*section',
                r'do not forget.*blacken.*id',
                r'avoid.*academic.*cheating',
                r'not.*taking.*part.*exam.*paper',
                r'failure.*incomplete.*grade',
                r'read.*questions.*carefully',
                r'blacken.*best.*answer',
                r'answer.*sheet',
                r'computerized.*answer.*sheet',
                r'subject.*deduction',
                
                # Part headers and instructions (excluding main part headers)
                r'instruction:.*',
                r'continue.*part',
                r'multiple.*choice.*questions',
                r'true.*false.*questions',
                r'matching.*questions',
                r'short.*questions',
                r'long.*questions',
                r'essay.*questions',
                
                # Common formatting
                r'^\d+\s*marks?\s*each',
                r'questions?\s*:\s*\d+\s*marks?',
                r'limit.*answers.*lines',
                r'use.*examples.*diagrams',
                r'answer.*booklet.*provided',
                r'clearly.*descriptively',
                
                # Generic headers/footers
                r'header', r'footer', r'page \d+',
                r'name:', r'date:', r'class:', r'section:', r'student id',
            ]
            
            def is_instruction_text(text: str) -> bool:
                """Check if text is a common exam instruction"""
                text_lower = text.lower().strip()
                
                # Skip empty or very short text
                if len(text_lower) < 5:
                    return True
                
                # IMPORTANT: Don't filter out Part headers - we need them for section detection
                if re.search(r'part\s+[ivx]+.*:', text_lower, re.IGNORECASE):
                    return False
                
                # Check against instruction patterns
                for pattern in instruction_patterns:
                    if re.search(pattern, text_lower, re.IGNORECASE):
                        return True
                
                # Skip lines that are mostly punctuation or formatting
                if len(re.sub(r'[^a-zA-Z0-9]', '', text_lower)) < 3:
                    return True
                
                return False
            
            # Extract all paragraphs with debug tracking
            all_paragraphs = []
            current_section = None  # Track current section type
            exam_content_started = False  # Flag to track when actual exam content begins
            section_order = []  # Track the order in which sections appear
            
            # First pass: Find where actual exam content starts
            exam_start_index = None
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    # Look for patterns that indicate actual exam content has started
                    if (re.search(r'Part\s+[IVX]+.*Multiple.*Choice.*Questions', text, re.IGNORECASE) or
                        re.search(r'Part\s+[IVX]+.*True.*False.*Questions', text, re.IGNORECASE) or
                        re.search(r'Part\s+[IVX]+.*Matching.*Questions', text, re.IGNORECASE) or
                        re.search(r'Part\s+[IVX]+.*Short.*Questions', text, re.IGNORECASE) or
                        re.search(r'Part\s+[IVX]+.*Long.*Questions', text, re.IGNORECASE) or
                        re.search(r'Part\s+[IVX]+.*Essay.*Questions', text, re.IGNORECASE)):
                        exam_start_index = i
                        break
            
            # If we couldn't find a clear start, look for first numbered question
            if exam_start_index is None:
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text and re.match(r'^1[\.\)\s]+', text):  # First question
                        exam_start_index = max(0, i - 5)  # Start a few lines before first question
                        break
            
            # Default to starting from paragraph 30 if no clear pattern found
            if exam_start_index is None:
                exam_start_index = 30
            
            logger.info(f"Detected exam content starting at paragraph index {exam_start_index}")
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # Skip metadata - only process paragraphs from exam start onwards
                if i < exam_start_index:
                    continue
                
                if text:
                    if not is_instruction_text(text):
                        # Detect section headers and track their order
                        section_detected = None
                        if re.search(r'Part\s+[IVX]+.*True.*False', text, re.IGNORECASE):
                            section_detected = 'true_false'
                        elif re.search(r'Part\s+[IVX]+.*Matching', text, re.IGNORECASE):
                            section_detected = 'matching'
                        elif re.search(r'Part\s+[IVX]+.*Short', text, re.IGNORECASE):
                            section_detected = 'short_answer'
                        elif re.search(r'Part\s+[IVX]+.*Long', text, re.IGNORECASE):
                            section_detected = 'long_answer'
                        elif re.search(r'Part\s+[IVX]+.*Multiple.*Choice', text, re.IGNORECASE):
                            section_detected = 'mcq'
                        
                        # Track section order if new section detected
                        if section_detected and section_detected not in section_order:
                            section_order.append(section_detected)
                            
                        if section_detected:
                            current_section = section_detected
                        
                        all_paragraphs.append((text, current_section))
            
            # Comprehensive question extraction for all types
            questions_by_type = {
                'mcq': [],
                'true_false': [],
                'matching': [],
                'short_answer': [],
                'long_answer': [],  # Changed from 'essay' to 'long_answer'
                'fill_blank': [],
                'other': []
            }
            
            # Keep track of question numbers within each section
            section_question_counts = {
                'mcq': 0,
                'true_false': 0,
                'matching': 0,
                'short_answer': 0,
                'long_answer': 0,
                'fill_blank': 0,
                'other': 0
            }
            
            i = 0
            while i < len(all_paragraphs):
                text, section_context = all_paragraphs[i]
                
                # Skip section headers and instruction lines
                if (re.search(r'Part\s+[IVX]+:', text, re.IGNORECASE) or
                    re.search(r'Continue to Part', text, re.IGNORECASE) or
                    re.search(r'Instruction:', text, re.IGNORECASE) or
                    re.search(r'Column [AB]', text, re.IGNORECASE)):
                    i += 1
                    continue
                
                # Detect question number patterns (more flexible)
                question_match = re.match(r'^(\d+)[\.\)\s:]+(.+)', text, re.IGNORECASE | re.MULTILINE)
                if not question_match:
                    # Also check for questions without explicit numbers in T/F sections
                    if (section_context == 'true_false' and 
                        re.search(r'\b(is|are|was|were|do|does|did|can|will|would|should)\b', text.lower()) and 
                        (text.endswith('?') or len(text.split()) > 3)):
                        # Create artificial question number for T/F without numbers
                        section_question_counts['true_false'] += 1
                        section_question_num = str(section_question_counts['true_false'])
                        question_content = text.strip()
                        complete_question = f"{section_question_num}. {question_content}"
                        question_type = 'true_false'
                        questions.append(complete_question)
                        questions_by_type['true_false'].append({
                            'number': section_question_num,
                            'content': complete_question,
                            'type': 'true_false'
                        })
                    i += 1
                    continue
                
                original_question_num = question_match.group(1)
                question_content = question_match.group(2).strip()
                
                # Build complete question by looking ahead for options/content
                complete_question = text
                question_type = 'other'
                j = i + 1
                
                # Look ahead for MCQ options (a, b, c, d)
                mcq_options = []
                while j < len(all_paragraphs) and j < i + 10:  # Look max 10 lines ahead
                    next_text, next_section = all_paragraphs[j]
                    
                    # Check if it's an MCQ option
                    mcq_match = re.match(r'^[a-eA-E][\.\)]\s*(.+)', next_text)
                    if mcq_match:
                        mcq_options.append(next_text)
                        complete_question += " " + next_text
                        j += 1
                    # Check if it's the start of next question
                    elif re.match(r'^\d+[\.\)\s]+', next_text):
                        break
                    # Check for section break
                    elif re.search(r'Continue to Part|Part\s+[IVX]+', next_text):
                        break
                    # Check for continuation of current question
                    elif len(next_text) > 5 and not re.match(r'^[A-Z][a-z]+\s*[IVX]+:', next_text):
                        complete_question += " " + next_text
                        j += 1
                    else:
                        break
                
                # Determine question type based on content, structure, and section context
                question_lower = question_content.lower()
                complete_lower = complete_question.lower()
                
                # Use section context as primary hint
                if section_context and section_context != 'mcq':
                    question_type = section_context if section_context != 'essay' else 'long_answer'
                elif mcq_options:
                    question_type = 'mcq'
                elif (re.search(r'\b(true|false)\b', question_lower) or 
                      re.search(r'\b(correct|incorrect)\b', question_lower) or
                      'T or F' in complete_question or
                      'True/False' in complete_question or
                      re.search(r'\bis\s+.*\?', complete_lower) or
                      re.search(r'\bare\s+.*\?', complete_lower) or
                      re.search(r'\bhas\s+.*\?', complete_lower) or
                      re.search(r'\bdoes\s+.*\?', complete_lower)):
                    question_type = 'true_false'
                elif (re.search(r'\b(match|matching|column)\b', question_lower) or
                      'Column A' in complete_question or 'Column B' in complete_question or
                      re.search(r'\b(something|term|appropriate)\b', complete_lower)):
                    question_type = 'matching'
                elif (re.search(r'\b(fill|blank|complete)\b', question_lower) or
                      '____' in complete_question or '___' in complete_question):
                    question_type = 'fill_blank'
                elif (re.search(r'\b(describe|explain|discuss|analyze|compare|essay|process)\b', question_lower) or
                      re.search(r'\b(write|composition|paragraph|clearly|descriptively)\b', question_lower) or
                      'examples or diagrams' in complete_lower or
                      len(question_content.split()) > 15):
                    question_type = 'long_answer'
                elif (re.search(r'\b(short|brief|list|define|name|what is)\b', question_lower) or
                      re.search(r'\b(ingredients|common|two|three)\b', question_lower) or
                      (len(question_content.split()) <= 15 and question_content.endswith('?'))):
                    question_type = 'short_answer'
                
                # Increment question count for this section and get section-specific number
                section_question_counts[question_type] += 1
                section_question_num = str(section_question_counts[question_type])
                
                # Clean and store the question
                clean_question = complete_question.strip()
                
                # For non-MCQ questions, add the section-specific number at the beginning if not already present
                if question_type != 'mcq' and not re.match(r'^\d+\.', clean_question):
                    clean_question = f"{section_question_num}. {clean_question}"
                
                if len(clean_question.split()) >= 3:  # Minimum word count
                    question_obj = {
                        'number': section_question_num,
                        'content': clean_question,
                        'type': question_type
                    }
                    
                    # For matching questions, try to separate individual questions and extract column information
                    if question_type == 'matching':
                        # Check if this is a combined matching question with multiple numbered items
                        individual_questions = self.parse_matching_questions(clean_question)
                        
                        if individual_questions:
                            # Process each individual matching question separately
                            for idx, individual_question in enumerate(individual_questions):
                                individual_question_obj = {
                                    'number': f"{section_question_num}.{idx + 1}" if len(individual_questions) > 1 else section_question_num,
                                    'content': individual_question,
                                    'type': question_type
                                }
                                
                                # Extract column A and column B parts for individual question
                                column_match = re.search(r'with\s+([A-H]\.\s*.+)', individual_question, re.IGNORECASE | re.DOTALL)
                                if column_match:
                                    column_b_part = column_match.group(1)
                                    column_a_part = individual_question.replace(f" with {column_b_part}", "").replace("Match:", "").strip()
                                    # Remove the question number from column_a if present
                                    column_a_part = re.sub(r'^\d+\.\s*', '', column_a_part)
                                    individual_question_obj["column_a"] = column_a_part
                                    individual_question_obj["column_b"] = column_b_part
                                
                                questions_by_type[question_type].append(individual_question_obj)
                                questions.append(individual_question)
                        else:
                            # Fallback to original logic for questions that don't match the pattern
                            column_match = re.search(r'with\s+([A-H]\.\s*.+)', clean_question, re.IGNORECASE | re.DOTALL)
                            if column_match:
                                column_b_part = column_match.group(1)
                                column_a_part = clean_question.replace(f" with {column_b_part}", "").replace("Match:", "").strip()
                                column_a_part = re.sub(r'^\d+\.\s*', '', column_a_part)
                                question_obj["column_a"] = column_a_part
                                question_obj["column_b"] = column_b_part
                            
                            questions_by_type[question_type].append(question_obj)
                            questions.append(clean_question)
                    else:
                        questions_by_type[question_type].append(question_obj)
                        questions.append(clean_question)
                
                # Move to next unprocessed paragraph
                i = j if j > i + 1 else i + 1
            
            # Also extract standalone questions (without numbers)
            for text, section_context in all_paragraphs:
                if (not re.match(r'^\d+[\.\)\s]+', text) and 
                    len(text) > 20 and
                    (text.endswith('?') or 
                     re.search(r'\b(what|where|when|why|how|which|who|is|are|do|does)\b', text.lower()))):
                    questions.append(text)
                    
                    # Use section context or detect type
                    q_type = section_context if section_context else 'other'
                    questions_by_type[q_type if q_type in questions_by_type else 'other'].append({
                        'number': 'unnumbered',
                        'content': text,
                        'type': q_type if q_type in questions_by_type else 'standalone'
                    })
            
            # EXTRACT QUESTIONS FROM TABLES
            
            # Build a map of document elements (paragraphs and tables) to determine section context
            document_elements = []
            
            # Add paragraphs with their positions
            for para in doc.paragraphs:
                document_elements.append(('paragraph', para.text.strip()))
            
            # Find table positions relative to paragraphs (approximate)
            table_positions = []
            for table_idx, table in enumerate(doc.tables):
                # Heuristic: tables usually appear after their section headers
                # Find most relevant section based on table content and position
                table_positions.append(table_idx)
            
            # Process tables, but apply exam content filtering
            tables_to_process = []
            for table_idx, table in enumerate(doc.tables):
                # Only process tables that appear after the exam content starts
                # Tables before the exam content start are likely metadata/course info
                if table_idx >= max(0, exam_start_index // 10):  # Approximate table position
                    tables_to_process.append((table_idx, table))
                    
            for table_idx, table in tables_to_process:
                if not table.rows:
                    continue
                    
                # Determine table type based on header row and content analysis
                header_row = table.rows[0] if table.rows else None
                if not header_row:
                    continue
                    
                header_texts = [cell.text.strip().lower() for cell in header_row.cells]
                table_type = 'unknown'
                
                # Analyze actual content to determine type more accurately
                sample_content = []
                sample_questions = []
                for row in table.rows[:5]:  # Check first 5 rows
                    if len(row.cells) >= 2:
                        cell_text = row.cells[1].text.strip()
                        if cell_text and len(cell_text) > 5:
                            sample_content.append(cell_text.lower())
                            sample_questions.append(cell_text)
                
                # Identify table type based on content patterns
                if len(header_texts) >= 2:
                    if 'column a' in header_texts[0] and 'column b' in header_texts[1]:
                        table_type = 'matching'
                    else:
                        # Analyze question content for type classification
                        true_false_indicators = 0
                        short_answer_indicators = 0
                        long_answer_indicators = 0
                        
                        for content in sample_content:
                            # True/False indicators
                            if (content.startswith('is ') or content.startswith('are ') or
                                content.startswith('do ') or content.startswith('does ') or
                                content.startswith('was ') or content.startswith('were ') or
                                content.startswith('has ') or content.startswith('have ') or
                                re.search(r'\b(is|are|was|were|has|have|can|will)\b.*\.$', content)):
                                true_false_indicators += 2
                            
                            # Short answer indicators
                            if (content.startswith('what ') or content.startswith('list ') or
                                content.startswith('name ') or content.startswith('define ') or
                                'what is' in content or 'common ingredients' in content or
                                content.endswith('?')):
                                short_answer_indicators += 2
                            
                            # Long answer indicators
                            if (any(word in content for word in ['describe', 'explain', 'discuss', 'process']) or
                                len(content.split()) > 12):
                                long_answer_indicators += 2
                        
                        # Classify based on strongest indicators
                        if long_answer_indicators > max(true_false_indicators, short_answer_indicators):
                            table_type = 'long_answer'
                        elif short_answer_indicators > true_false_indicators:
                            table_type = 'short_answer'
                        elif true_false_indicators > 0:
                            table_type = 'true_false'
                        else:
                            # Default fallback based on average question length
                            avg_length = sum(len(content.split()) for content in sample_content) / len(sample_content) if sample_content else 0
                            if avg_length > 15:
                                table_type = 'long_answer'
                            else:
                                table_type = 'short_answer'
                
                # Extract questions based on table type
                if table_type == 'matching':
                    # Column A and Column B format
                    for row_idx, row in enumerate(table.rows[1:], 1):  # Skip header row
                        if len(row.cells) >= 2:
                            col_a = row.cells[0].text.strip()
                            col_b = row.cells[1].text.strip()
                            
                            if col_a and col_b:
                                # Check if col_a contains multiple numbered items (like "1. Something sour. 2. Something sweet.")
                                combined_content = f"Match: {col_a} with {col_b}"
                                individual_questions = self.parse_matching_questions(combined_content)
                                
                                if individual_questions:
                                    # Process each individual matching question separately
                                    for idx, clean_question_content in enumerate(individual_questions):
                                        section_question_counts['matching'] += 1
                                        matching_number = str(section_question_counts['matching'])
                                        
                                        # Store clean content for display
                                        questions.append(clean_question_content)
                                        questions_by_type['matching'].append({
                                            'number': matching_number,
                                            'content': clean_question_content,  # Clean content without "Match:" prefix
                                            'type': 'matching',
                                            'full_content': combined_content,  # Keep original for reference
                                            'column_a': col_a,
                                            'column_b': col_b
                                        })
                                else:
                                    # Fallback to original logic for simple matching questions
                                    section_question_counts['matching'] += 1
                                    matching_number = str(section_question_counts['matching'])
                                    matching_question = f"Match: {col_a} with {col_b}"
                                    questions.append(matching_question)
                                    questions_by_type['matching'].append({
                                        'number': matching_number,
                                        'content': matching_question,
                                        'type': 'matching',
                                        'column_a': col_a,
                                        'column_b': col_b
                                    })
                
                elif table_type in ['true_false', 'short_answer', 'long_answer']:
                    # Number in col1, question in col2 format
                    for row_idx, row in enumerate(table.rows):
                        if len(row.cells) >= 2:
                            number_cell = row.cells[0].text.strip()
                            question_cell = row.cells[1].text.strip()
                            
                            # Skip header row and empty rows
                            if (question_cell and 
                                not question_cell.lower().startswith('column') and
                                len(question_cell) > 5 and
                                not question_cell.lower() in ['course title:', 'lecturer:', 'date:']):
                                
                                # Use number from first column or generate one
                                question_num = number_cell if number_cell.isdigit() else f'table_{table_idx}_row_{row_idx}'
                                
                                questions.append(question_cell)
                                questions_by_type[table_type].append({
                                    'number': question_num,
                                    'content': question_cell,
                                    'type': table_type
                                })
            
            # Extract additional patterns for completeness
            mcq_options = []
            matching_pairs = []
            true_false_statements = []
            
            for text, section_context in all_paragraphs:
                # MCQ options
                if re.match(r'^[a-eA-E][\.\)]\s*\w+', text):
                    mcq_options.append(text)
                
                # Matching pairs (Column A/B format or single letter format)
                if (re.search(r'^[A-Z]\.\s*\w+', text) or 
                    re.search(r'^\s*[A-H]\.\s*\w+', text)):
                    matching_pairs.append(text)
                
                # True/False statements
                if re.search(r'\b(true|false|T|F)\b', text, re.IGNORECASE) and len(text.split()) > 2:
                    true_false_statements.append(text)
            
            logger.info(f"Extracted {len(questions)} questions from {os.path.basename(file_path)}")
            
            # Prepare comprehensive extraction data
            # Create dynamic part mapping based on section order
            dynamic_part_mapping = {}
            roman_numerals = ['I', 'II', 'III', 'IV', 'V']
            
            for idx, section_type in enumerate(section_order):
                if idx < len(roman_numerals):
                    dynamic_part_mapping[section_type] = f"Part {roman_numerals[idx]}"
                else:
                    dynamic_part_mapping[section_type] = f"Part {idx + 1}"
            
            extraction_data = {
                "file_info": {
                    "filename": os.path.basename(file_path),
                    "extraction_time": datetime.now().isoformat()
                },
                "questions": questions,
                "questions_by_type": questions_by_type,
                "dynamic_part_mapping": dynamic_part_mapping,
                "question_statistics": {
                    "total_questions": len(questions),
                    "mcq_count": len(questions_by_type.get('mcq', [])),
                    "true_false_count": len(questions_by_type.get('true_false', [])),
                    "matching_count": len(questions_by_type.get('matching', [])),
                    "short_answer_count": len(questions_by_type.get('short_answer', [])),
                    "long_answer_count": len(questions_by_type.get('long_answer', [])),
                    "fill_blank_count": len(questions_by_type.get('fill_blank', [])),
                    "other_count": len(questions_by_type.get('other', []))
                }
            }
            
            return extraction_data
            
        except Exception as e:
            logger.error(f"Error extracting questions from {file_path}: {str(e)}")
            
            # Create minimal error data structure
            error_data = {
                "file_info": {
                    "filename": os.path.basename(file_path),
                    "full_path": file_path,
                    "extraction_time": datetime.now().isoformat(),
                    "error_occurred": True
                },
                "questions": [],
                "questions_by_type": {
                    'mcq': [], 'true_false': [], 'matching': [],
                    'short_answer': [], 'long_answer': [], 'fill_blank': [], 'other': []
                },
                "question_statistics": {
                    "total_questions": 0,
                    "mcq_count": 0, "true_false_count": 0, "matching_count": 0,
                    "short_answer_count": 0, "long_answer_count": 0, "fill_blank_count": 0, "other_count": 0
                },
                "error_details": {
                    "error_message": str(e),
                    "error_type": type(e).__name__
                }
            }
            
            return error_data
    
    def parse_matching_questions(self, matching_text: str) -> List[str]:
        """Parse a combined matching question into individual numbered questions"""
        try:
            # Remove the "Match:" prefix and clean up
            text_content = matching_text.replace("Match:", "").strip()
            
            # First, separate the "with" part (Column B options) if present
            with_match = re.search(r'\s+with\s+([A-H]\.\s*.+)', text_content, re.IGNORECASE | re.DOTALL)
            column_b_part = ""
            column_a_content = text_content
            
            if with_match:
                column_b_part = with_match.group(1).strip()
                column_a_content = text_content.replace(with_match.group(0), "").strip()
            
            # Pattern to match numbered items in Column A like "1. Something sour." or "2. Something sweet."
            item_pattern = r'(\d+)[\.\)\s]+([^0-9]+?)(?=\s*\d+[\.\)\s]+|$)'
            
            matches = re.findall(item_pattern, column_a_content, re.DOTALL)
            
            if len(matches) < 2:  # If we can't find at least 2 numbered items, return original
                return []
            
            individual_questions = []
            for number, content in matches:
                # Clean up the content - remove extra whitespace and trailing punctuation
                clean_content = re.sub(r'\s+', ' ', content.strip())
                clean_content = clean_content.rstrip('.,')  # Remove trailing punctuation
                
                if clean_content:  # Only add non-empty content
                    # Create individual question with CLEAN content (no "Match: number." prefix, no options)
                    # Just the clean content with proper punctuation
                    if not clean_content.endswith('.'):
                        clean_content += '.'
                    
                    individual_questions.append(clean_content)
            
            return individual_questions if len(individual_questions) >= 2 else []
            
        except Exception as e:
            logger.error(f"Error parsing matching questions: {str(e)}")
            return []

    def calculate_tfidf_similarity(self, text1: str, text2: str) -> float:
        """Calculate TF-IDF similarity between two texts"""
        try:
            documents = [text1, text2]
            tfidf_matrix = self.tfidf_vectorizer.fit_transform(documents)
            similarity_matrix = cosine_similarity(tfidf_matrix)
            return float(similarity_matrix[0][1])
        except Exception as e:
            logger.error(f"Error calculating TF-IDF similarity: {str(e)}")
            return 0.0
    
    def normalize_text(self, text: str) -> str:
        """Normalize text for comparison"""
        if not text or not isinstance(text, str):
            return ""
        
        # Convert to lowercase
        text = text.lower().strip()
        
        # Remove extra whitespace and newlines
        text = re.sub(r'\s+', ' ', text)
        
        # Remove punctuation but keep question marks
        text = re.sub(r'[^\w\s\?]', ' ', text)
        
        # Remove extra spaces
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    def calculate_enhanced_similarity(self, text1: str, text2: str) -> float:
        """Simpler similarity: TF-IDF + Jaccard"""
        try:
            norm1 = self.normalize_text(text1)
            norm2 = self.normalize_text(text2)

            # TF-IDF similarity
            tfidf_similarity = self.calculate_tfidf_similarity(norm1, norm2)

            # Jaccard similarity
            words1, words2 = set(norm1.split()), set(norm2.split())
            jaccard_similarity = (
                len(words1 & words2) / len(words1 | words2)
                if words1 and words2 else 0.0
            )

            # Average (equal weights)
            return (tfidf_similarity + jaccard_similarity) / 2
            
        except Exception as e:
            logger.error(f"Error calculating enhanced similarity: {str(e)}")
            return 0.0
    
    def compare_documents(self, doc1_content: Dict, doc2_content: Dict) -> Dict[str, Any]:
        """Compare two documents using TF-IDF and semantic analysis"""
        questions1 = " ".join(doc1_content.get("questions", []))
        questions2 = " ".join(doc2_content.get("questions", []))
        
        if not questions1 or not questions2:
            return {
                "similarity_score": 0.0,
                "method_used": "none",
                "details": "One or both documents have no extractable questions",
                "processing_time": 0.0
            }
        
        start_time = time.time()
        
        # Step 1: TF-IDF Similarity
        tfidf_score = self.calculate_tfidf_similarity(questions1, questions2)
        
        method_used = "tfidf"
        final_score = tfidf_score
        
        # Step 2: Semantic Analysis (if TF-IDF < threshold)
        if tfidf_score < self.similarity_threshold:
            semantic_score = self.calculate_enhanced_similarity(questions1, questions2)
            final_score = max(tfidf_score, semantic_score)
            method_used = "tfidf_enhanced"
        
        processing_time = time.time() - start_time
        
        # Find potentially matching questions
        matching_questions = self.find_matching_questions(doc1_content, doc2_content)
        
        return {
            "similarity_score": round(final_score, 3),
            "tfidf_score": round(tfidf_score, 3),
            "semantic_score": round(self.calculate_enhanced_similarity(questions1, questions2), 3) if tfidf_score < self.similarity_threshold else None,
            "method_used": method_used,
            "matching_questions": matching_questions,
            "processing_time": round(processing_time, 3)
        }
    
    def find_matching_questions(self, doc1_content: Dict, doc2_content: Dict) -> List[Dict]:
        """Find potentially matching questions between documents using dynamic part detection"""
        matches = []
        
        # Get structured questions by type from both documents
        questions1_by_type = doc1_content.get("questions_by_type", {})
        questions2_by_type = doc2_content.get("questions_by_type", {})
        
        # Get dynamic part mappings from both documents
        part_mapping1 = doc1_content.get("dynamic_part_mapping", {})
        part_mapping2 = doc2_content.get("dynamic_part_mapping", {})
        
        # Create flat lists with part-specific indexing information
        structured_questions1 = []
        structured_questions2 = []
        
        # Build structured question lists for document 1
        for q_type, questions in questions1_by_type.items():
            part_name = part_mapping1.get(q_type, q_type.title())
            for q_obj in questions:
                # Get question content and clean MCQ display
                content = q_obj.get('content', '')
                display_content = content
                
                # For MCQ questions, remove first 2 characters (number and period/space) for display
                if q_type == 'mcq' and len(content) > 2:
                    # Remove pattern like "1. " or "1) " from beginning
                    display_content = re.sub(r'^\d+[\.\)\s]+', '', content).strip()
                
                structured_questions1.append({
                    'content': content,  # Keep original for similarity comparison
                    'display_content': display_content,  # Clean version for display
                    'number': q_obj.get('number', 'N/A'),
                    'type': q_type,
                    'part': part_name,
                    'display_index': f"{part_name}: Q{q_obj.get('number', 'N/A')}"
                })
        
        # Build structured question lists for document 2  
        for q_type, questions in questions2_by_type.items():
            part_name = part_mapping2.get(q_type, q_type.title())
            for q_obj in questions:
                # Get question content and clean MCQ display
                content = q_obj.get('content', '')
                display_content = content
                
                # For MCQ questions, remove first 2 characters (number and period/space) for display
                if q_type == 'mcq' and len(content) > 2:
                    # Remove pattern like "1. " or "1) " from beginning
                    display_content = re.sub(r'^\d+[\.\)\s]+', '', content).strip()
                
                structured_questions2.append({
                    'content': content,  # Keep original for similarity comparison
                    'display_content': display_content,  # Clean version for display
                    'number': q_obj.get('number', 'N/A'),
                    'type': q_type,
                    'part': part_name,
                    'display_index': f"{part_name}: Q{q_obj.get('number', 'N/A')}"
                })
        
        # Compare questions using structured data
        for i, q1_struct in enumerate(structured_questions1):
            for j, q2_struct in enumerate(structured_questions2):
                # Use enhanced similarity calculation for better accuracy
                similarity = self.calculate_enhanced_similarity(q1_struct['content'], q2_struct['content'])
                
                # Debug logging to understand similarity scores
                if similarity > 0.2:  # Log scores above 0.2 to see what we're getting
                    logger.info(f"Question similarity: {similarity:.3f} - Q1: '{q1_struct['content'][:50]}...' vs Q2: '{q2_struct['content'][:50]}...'")
                
                if similarity > 0.25:  # Lowered threshold from 0.6 to 0.25 for enhanced similarity
                    # Use display content for showing to user (cleaned MCQ content)
                    q1_display_text = q1_struct['display_content']
                    q2_display_text = q2_struct['display_content']
                    
                    matches.append({
                        "question1_index": q1_struct['display_index'],  # Part-specific index
                        "question2_index": q2_struct['display_index'],  # Part-specific index
                        "question1": q1_display_text[:100] + "..." if len(q1_display_text) > 100 else q1_display_text,
                        "question2": q2_display_text[:100] + "..." if len(q2_display_text) > 100 else q2_display_text,
                        "question1_type": q1_struct['type'],
                        "question2_type": q2_struct['type'], 
                        "similarity": round(similarity, 3)
                    })
        
        # Sort by similarity score (highest first)
        matches.sort(key=lambda x: x["similarity"], reverse=True)
        return matches[:20]  # Return top 20 matches (increased from 10 to catch more partial matches)
    
    def analyze_session(self, session_id: str) -> Dict[str, Any]:
        """Analyze all documents in a session and create similarity matrix"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        uploaded_files_path = os.path.join(session_path, "uploaded_files")
        extracted_content_path = os.path.join(session_path, "extracted_content")
        results_path = os.path.join(session_path, "analysis_results")
        
        if not os.path.exists(uploaded_files_path):
            raise ValueError(f"Session {session_id} not found")
        
        # Get all uploaded files
        uploaded_files = [f for f in os.listdir(uploaded_files_path) if f.endswith('.docx')]
        
        if len(uploaded_files) < 2:
            raise ValueError("At least 2 documents required for comparison")
        
        # Extract content from all documents
        extracted_content = {}
        
        for filename in uploaded_files:
            file_path = os.path.join(uploaded_files_path, filename)
            logger.info(f"Extracting content from: {filename}")
            try:
                content = self.extract_questions_from_docx(file_path)
                extracted_content[filename] = content
                logger.info(f"Successfully extracted {content.get('question_statistics', {}).get('total_questions', 0)} questions from {filename}")
            except Exception as e:
                logger.error(f"Failed to extract content from {filename}: {str(e)}")
                # Create empty content structure for failed extractions
                extracted_content[filename] = {
                    "questions": [],
                    "questions_by_type": {},
                    "question_statistics": {"total_questions": 0}
                }
        
        # Create similarity matrix
        detailed_comparisons = {}
        
        # Initialize matrix with filenames
        filenames = list(extracted_content.keys())
        matrix_data = {
            "filenames": filenames,
            "matrix": []
        }
        
        # Calculate similarities
        for i, file1 in enumerate(filenames):
            row = []
            for j, file2 in enumerate(filenames):
                if i == j:
                    # Same file - similarity = 1.0
                    row.append(1.0)
                elif f"{file1}_vs_{file2}" in detailed_comparisons or f"{file2}_vs_{file1}" in detailed_comparisons:
                    # Already calculated (symmetric)
                    existing_key = f"{file2}_vs_{file1}" if f"{file2}_vs_{file1}" in detailed_comparisons else f"{file1}_vs_{file2}"
                    row.append(detailed_comparisons[existing_key]["similarity_score"])
                else:
                    # Calculate new comparison
                    comparison = self.compare_documents(
                        extracted_content[file1], 
                        extracted_content[file2]
                    )
                    detailed_comparisons[f"{file1}_vs_{file2}"] = comparison
                    row.append(comparison["similarity_score"])
            
            matrix_data["matrix"].append(row)
        
        # Prepare comprehensive results data
        comprehensive_results = {
            "analysis_metadata": {
                "session_id": session_id,
                "analysis_completion_time": datetime.now().isoformat(),
                "total_files_analyzed": len(filenames),
                "total_comparisons_made": len(detailed_comparisons),
                "analysis_method": "tfidf_with_semantic_fallback"
            },
            "file_statistics": {
                filename: {
                    "total_questions": extracted_content[filename].get("question_statistics", {}).get("total_questions", 0),
                    "question_breakdown": extracted_content[filename].get("question_statistics", {}),
                    "processing_metadata": extracted_content[filename].get("processing_metadata", {})
                }
                for filename in filenames
            },
            "similarity_matrix": matrix_data,
            "detailed_comparisons": detailed_comparisons
        }
        
        # Save individual components
        matrix_file = os.path.join(results_path, "similarity_matrix.json")
        with open(matrix_file, 'w', encoding='utf-8') as f:
            json.dump(matrix_data, f, indent=2, ensure_ascii=False)
        
        details_file = os.path.join(results_path, "detailed_comparisons.json")
        with open(details_file, 'w', encoding='utf-8') as f:
            json.dump(detailed_comparisons, f, indent=2, ensure_ascii=False)
        
        # Save comprehensive results
        comprehensive_file = os.path.join(results_path, "comprehensive_analysis_results.json")
        with open(comprehensive_file, 'w', encoding='utf-8') as f:
            json.dump(comprehensive_results, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved comprehensive analysis results to: {comprehensive_file}")
        
        # Update session info
        session_info_path = os.path.join(session_path, "session_info.json")
        with open(session_info_path, 'r') as f:
            session_info = json.load(f)
        
        session_info["analysis_completed"] = True
        session_info["analysis_time"] = datetime.now().isoformat()
        session_info["total_comparisons"] = len(detailed_comparisons)
        
        with open(session_info_path, 'w') as f:
            json.dump(session_info, f, indent=2)
        
        logger.info(f"Analysis completed for session {session_id}")
        
        return {
            "session_id": session_id,
            "similarity_matrix": matrix_data,
            "detailed_comparisons": detailed_comparisons,
            "files_analyzed": len(filenames),
            "total_comparisons": len(detailed_comparisons)
        }
    
    def cleanup_session(self, session_id: str) -> bool:
        """Remove a specific session and all its files"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        
        if os.path.exists(session_path):
            try:
                shutil.rmtree(session_path)
                logger.info(f"Cleaned up session: {session_id}")
                return True
            except Exception as e:
                logger.error(f"Error cleaning up session {session_id}: {str(e)}")
                return False
        
        return False
    
    def cleanup_old_sessions(self, max_age_hours: int = 2) -> int:
        """Remove sessions older than specified hours and empty sessions"""
        if not os.path.exists(self.temp_base_dir):
            return 0
        
        cutoff_time = time.time() - (max_age_hours * 3600)
        cleaned_count = 0
        
        for session_folder in os.listdir(self.temp_base_dir):
            if session_folder.startswith("similarity_"):
                session_path = os.path.join(self.temp_base_dir, session_folder)
                if os.path.isdir(session_path):
                    should_cleanup = False
                    
                    # Check if session is too old
                    if os.path.getctime(session_path) < cutoff_time:
                        should_cleanup = True
                        logger.info(f"Session {session_folder} is too old, cleaning up")
                    
                    # Check if session is empty (no uploaded files)
                    uploaded_files_path = os.path.join(session_path, "uploaded_files")
                    if os.path.exists(uploaded_files_path):
                        uploaded_files = [f for f in os.listdir(uploaded_files_path) if f.endswith('.docx')]
                        if len(uploaded_files) == 0:
                            # Also check if session is older than 5 minutes (to avoid cleaning up very recent empty sessions)
                            session_age_minutes = (time.time() - os.path.getctime(session_path)) / 60
                            if session_age_minutes > 5:
                                should_cleanup = True
                                logger.info(f"Session {session_folder} is empty and older than 5 minutes, cleaning up")
                    
                    if should_cleanup and self.cleanup_session(session_folder):
                        cleaned_count += 1
        
        logger.info(f"Cleaned up {cleaned_count} old/empty sessions")
        return cleaned_count
    
    def save_file_upload_info(self, session_id: str, uploaded_files: List[Dict[str, Any]]) -> bool:
        """Save information about uploaded files to the session"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        
        if not os.path.exists(session_path):
            logger.error(f"Session {session_id} not found")
            return False
        
        upload_info = {
            "upload_timestamp": datetime.now().isoformat(),
            "total_files": len(uploaded_files),
            "files": uploaded_files,
            "upload_metadata": {
                "session_id": session_id,
                "upload_successful": True
            }
        }
        
        upload_info_file = os.path.join(session_path, "file_upload_info.json")
        
        try:
            with open(upload_info_file, 'w', encoding='utf-8') as f:
                json.dump(upload_info, f, indent=2, ensure_ascii=False)
            
            # Also update session info with file list
            session_info_path = os.path.join(session_path, "session_info.json")
            if os.path.exists(session_info_path):
                with open(session_info_path, 'r', encoding='utf-8') as f:
                    session_info = json.load(f)
                
                session_info["files_uploaded"] = [f.get("filename", "unknown") for f in uploaded_files]
                session_info["last_file_upload"] = datetime.now().isoformat()
                
                with open(session_info_path, 'w', encoding='utf-8') as f:
                    json.dump(session_info, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Saved file upload info for {len(uploaded_files)} files in session {session_id}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to save file upload info: {str(e)}")
            return False
    
    def get_session_results(self, session_id: str) -> Dict[str, Any]:
        """Get analysis results for a session"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        results_path = os.path.join(session_path, "analysis_results")
        
        if not os.path.exists(results_path):
            raise ValueError(f"Results not found for session {session_id}")
        
        # Load similarity matrix
        matrix_file = os.path.join(results_path, "similarity_matrix.json")
        with open(matrix_file, 'r') as f:
            similarity_matrix = json.load(f)
        
        # Load detailed comparisons
        details_file = os.path.join(results_path, "detailed_comparisons.json")
        with open(details_file, 'r') as f:
            detailed_comparisons = json.load(f)
        
        # Load session info
        session_info_path = os.path.join(session_path, "session_info.json")
        with open(session_info_path, 'r') as f:
            session_info = json.load(f)
        
        return {
            "session_info": session_info,
            "similarity_matrix": similarity_matrix,
            "detailed_comparisons": detailed_comparisons
        }
