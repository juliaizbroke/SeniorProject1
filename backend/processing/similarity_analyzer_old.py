"""
Similarity Analysis System for Word Documents
Implements TF-IDF + Semantic similarity checking for exam papers
"""

import os
import re
import json
import time
import logging
import tempfile
from datetime import datetime
from typing import Dict, List, Any, Tuple
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
        
        # Create session directory structure
        session_path = os.path.join(self.temp_base_dir, session_id)
        os.makedirs(session_path, exist_ok=True)
        os.makedirs(os.path.join(session_path, "uploaded_files"), exist_ok=True)
        os.makedirs(os.path.join(session_path, "extracted_content"), exist_ok=True)
        os.makedirs(os.path.join(session_path, "analysis_results"), exist_ok=True)
        
        # Store session metadata
        session_info = {
            "session_id": session_id,
            "created_at": datetime.now().isoformat(),
            "status": "active",
            "uploaded_files": [],
            "analysis_complete": False
        }
        
        with open(os.path.join(session_path, "session_info.json"), 'w') as f:
            json.dump(session_info, f, indent=2)
        
        logger.info(f"Created similarity session: {session_id}")
        return session_id
    
    def extract_questions_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract questions from Word document using section-based detection"""
        
        def is_instruction_text(text: str) -> bool:
            """Check if text is an instruction to be filtered out"""
            text_lower = text.lower()
            instruction_patterns = [
                r'answer.*questions.*using.*pencil',
                r'write.*id.*name.*section',
                r'do not forget.*blacken.*id',
                r'avoid.*academic.*cheating',
                r'not.*taking.*part.*exam.*paper',
                r'failure.*incomplete.*grade',
                r'read.*questions.*carefully',
                r'blacken.*best.*answer',
                r'answer.*sheet',
                r'2b.*pencil',
                r'exam.*instructions'
            ]
            
            for pattern in instruction_patterns:
                if re.search(pattern, text_lower):
                    return True
            return False
        
        try:
            doc = Document(file_path)
            all_paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Section-based question type detection
            questions = []
            question_types = {
                'multiple_choice': [],
                'true_false': [],
                'matching': [], 
                'short_answer': [],
                'long_answer': [],
                'essay': [],
                'other': []
            }
            
            current_section_type = None
            i = 0
            
            while i < len(all_paragraphs):
                text = all_paragraphs[i]
                
                # Skip general instruction paragraphs
                if is_instruction_text(text):
                    i += 1
                    continue
                
                # Check for Part headers (Part I, Part II, etc.)
                part_match = re.match(r'^Part\s+[IVX]+[:.]\s*(.*)', text, re.IGNORECASE)
                if part_match:
                    section_title = part_match.group(1).lower()
                    
                    # Determine question type from section title
                    if 'multiple choice' in section_title or 'mcq' in section_title:
                        current_section_type = 'multiple_choice'
                    elif 'true' in section_title and 'false' in section_title:
                        current_section_type = 'true_false'
                    elif 'matching' in section_title or 'match' in section_title:
                        current_section_type = 'matching'
                    elif 'short' in section_title and 'question' in section_title:
                        current_section_type = 'short_answer'
                    elif 'long' in section_title and 'question' in section_title:
                        current_section_type = 'long_answer'
                    elif 'essay' in section_title:
                        current_section_type = 'essay'
                    else:
                        current_section_type = 'other'
                    
                    logger.info(f"Found section: {section_title} -> {current_section_type}")
                    i += 1
                    continue
                
                # Skip section instructions
                if text.lower().startswith('instruction:'):
                    i += 1
                    continue
                
                # Skip column headers for matching
                if text.lower() in ['column a', 'column b']:
                    i += 1
                    continue
                
                # Detect actual questions
                is_question = False
                
                # Question patterns - numbered questions and lettered options
                if re.match(r'^\d+[\.\s]', text):  # Starts with number
                    if len(text.split()) >= 3 and not is_instruction_text(text):
                        is_question = True
                elif re.match(r'^[a-h]\.\s+', text, re.IGNORECASE):  # MCQ/Matching options
                    if len(text.split()) >= 2 and not is_instruction_text(text):
                        # Only treat as question if we're in matching section (for Column A/B items)
                        # or if it's clearly an MCQ option following a question
                        if current_section_type == 'matching':
                            is_question = True
                
                if is_question and current_section_type:
                    question_content = text
                    
                    # For MCQ, collect options that follow
                    if current_section_type == 'multiple_choice':
                        j = i + 1
                        while j < len(all_paragraphs) and j < i + 8:
                            next_text = all_paragraphs[j]
                            if re.match(r'^[a-h]\.\s+', next_text, re.IGNORECASE):
                                question_content += " " + next_text
                                j += 1
                            else:
                                break
                        i = j - 1
                    
                    # Add question to appropriate category
                    questions.append(question_content.strip())
                    question_types[current_section_type].append(question_content.strip())
                
                i += 1
            
            # Extract MCQ options for counting
            mcq_options = []
            for mcq_question in question_types['multiple_choice']:
                # Split by spaces and find options manually
                words = mcq_question.split()
                for idx, word in enumerate(words):
                    if re.match(r'^[a-h]\.$', word, re.IGNORECASE):
                        # Found an option letter, get the next word(s)
                        if idx + 1 < len(words):
                            option_text = word + ' ' + words[idx + 1]
                            mcq_options.append(option_text)
            
            logger.info(f"Extracted {len(questions)} questions from {os.path.basename(file_path)}")
            logger.info(f"Question types: MCQ={len(question_types['multiple_choice'])}, T/F={len(question_types['true_false'])}, "
                       f"Matching={len(question_types['matching'])}, Short={len(question_types['short_answer'])}, "
                       f"Long={len(question_types['long_answer'])}, Essay={len(question_types['essay'])}, "
                       f"Other={len(question_types['other'])}")
            
            return {
                "questions": questions,
                "question_types": question_types,
                "mcq_options": mcq_options,
                "total_questions": len(questions),
                "extraction_time": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error extracting questions from {file_path}: {str(e)}")
            return {
                "questions": [],
                "question_types": {
                    'multiple_choice': [], 'true_false': [], 'matching': [],
                    'short_answer': [], 'long_answer': [], 'essay': [], 'other': []
                },
                "mcq_options": [],
                "total_questions": 0,
                "extraction_time": datetime.now().isoformat()
            }
import time
import shutil
import numpy as np
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Any
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import re
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentSimilarityAnalyzer:
    def __init__(self, temp_base_dir: str = "similarityCheck/similarity_temp"):
        self.temp_base_dir = temp_base_dir
        self.similarity_threshold = 0.7
        self.tfidf_vectorizer = TfidfVectorizer(
            stop_words='english',
            max_features=5000,
            ngram_range=(1, 2),
            lowercase=True,
            strip_accents='unicode'
        )
        
        # Ensure temp directory exists
        os.makedirs(temp_base_dir, exist_ok=True)
    
    def create_session(self) -> str:
        """Create a new similarity analysis session"""
        timestamp = int(time.time())
        random_suffix = os.urandom(4).hex()
        session_id = f"similarity_{timestamp}_{random_suffix}"
        
        session_path = os.path.join(self.temp_base_dir, session_id)
        os.makedirs(session_path, exist_ok=True)
        os.makedirs(os.path.join(session_path, "uploaded_files"), exist_ok=True)
        os.makedirs(os.path.join(session_path, "extracted_content"), exist_ok=True)
        os.makedirs(os.path.join(session_path, "analysis_results"), exist_ok=True)
        
        # Create session info
        session_info = {
            "session_id": session_id,
            "created_at": datetime.now().isoformat(),
            "files_uploaded": [],
            "analysis_completed": False,
            "expires_at": (datetime.now() + timedelta(hours=2)).isoformat()
        }
        
        with open(os.path.join(session_path, "session_info.json"), 'w') as f:
            json.dump(session_info, f, indent=2)
        
        logger.info(f"Created similarity session: {session_id}")
        return session_id
    
    def extract_questions_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract questions and answers from Word document"""
        try:
            doc = Document(file_path)
            questions = []
            current_question = ""
            
            # Common exam instructions to ignore
            instruction_patterns = [
                # Basic instructions
                r'answer.*questions.*using.*pencil',
                r'write.*id.*name.*section',
                r'do not forget.*blacken.*id',
                r'avoid.*academic.*cheating',
                r'not.*taking.*part.*exam.*paper',
                r'failure.*incomplete.*grade',
                r'read.*questions.*carefully',
                r'blacken.*best.*answer',
                r'answer.*sheet',
                r'computerized.*answer.*sheet',
                r'subject.*deduction',
                
                # Part headers and instructions
                r'part\s+[ivx]+:.*questions',
                r'instruction:.*',
                r'continue.*part',
                r'multiple.*choice.*questions',
                r'true.*false.*questions',
                r'matching.*questions',
                r'short.*questions',
                r'long.*questions',
                r'essay.*questions',
                
                # Common formatting
                r'^\d+\s*marks?\s*each',
                r'questions?\s*:\s*\d+\s*marks?',
                r'limit.*answers.*lines',
                r'use.*examples.*diagrams',
                r'answer.*booklet.*provided',
                r'clearly.*descriptively',
                
                # Generic headers/footers
                r'header', r'footer', r'page \d+',
                r'name:', r'date:', r'class:', r'section:', r'student id',
            ]
            
            def is_instruction_text(text: str) -> bool:
                """Check if text is a common exam instruction"""
                text_lower = text.lower().strip()
                
                # Skip empty or very short text
                if len(text_lower) < 5:
                    return True
                
                # Check against instruction patterns
                for pattern in instruction_patterns:
                    if re.search(pattern, text_lower, re.IGNORECASE):
                        return True
                
                # Skip lines that are mostly punctuation or formatting
                if len(re.sub(r'[^a-zA-Z0-9]', '', text_lower)) < 3:
                    return True
                
                return False
            
            # Process all paragraphs to extract different question types
            all_paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Filter out instruction paragraphs first
            content_paragraphs = []
            for text in all_paragraphs:
                if not is_instruction_text(text):
                    content_paragraphs.append(text)
            
            # Now extract questions from the filtered content
            i = 0
            while i < len(content_paragraphs):
                text = content_paragraphs[i]
                
                # Question type detection patterns
                question_patterns = [
                    # Multiple Choice Questions
                    r'^\d+\.\s+.*[?.]',                    # 1. What is...?
                    r'^Q\d+[.:]?\s+.*',                    # Q1: Which of...
                    r'^Question\s+\d+[.:]?\s+.*',          # Question 1: What...
                    
                    # True/False Questions  
                    r'^\d+\s+.*[.?]$',                     # 1   Is a stove used...?
                    r'^T/F[.:]?\s+.*',                     # T/F: Statement here
                    
                    # Matching Questions (Column headers and items)
                    r'^\d+\.\s+.*\.',                      # 1. Something sour.
                    r'^[A-Z]\.\s+.*',                      # A. Lime
                    
                    # Short/Long Answer Questions
                    r'^[IV]+\.\s+.*[?.]',                  # IV. Question text
                    r'^\d+\s+.*[?.]',                      # 1 What is an apple?
                    
                    # Essay Questions
                    r'^.*describe.*[?.]',                  # Describe your favourite...
                    r'^.*explain.*[?.]',                   # Explain the process...
                    r'^.*discuss.*[?.]',                   # Discuss the impact...
                    r'^.*analyze.*[?.]',                   # Analyze the following...
                    
                    # General question indicators
                    r'.*\b(what|where|when|why|how|which|who|describe|explain|list|name|identify)\b.*[?.]'
                ]
                
                # Check if current text matches any question pattern
                is_question = False
                for pattern in question_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        is_question = True
                        break
                
                # Additional content-based checks for questions
                if not is_question:
                    # Check for question words and question marks
                    if (len(text) > 10 and 
                        (text.endswith('?') or 
                         re.search(r'\b(what|where|when|why|how|which|who|describe|explain|list|identify)\b', text, re.IGNORECASE))):
                        is_question = True
                    
                    # Check for True/False format
                    elif re.search(r'(is|are|does|do|can|will|should|has|have).*[?.]', text, re.IGNORECASE):
                        is_question = True
                
                if is_question:
                    question_content = text
                    
                    # Look ahead to collect related content (options, etc.)
                    j = i + 1
                    while j < len(content_paragraphs) and j < i + 8:  # Look at next 7 lines max
                        next_text = content_paragraphs[j]
                        
                        # Check if it's part of current question (options, continuation, etc.)
                        is_related = (
                            # MCQ options (a, b, c, d format only)
                            re.match(r'^[a-d]\)|^[a-d]\.|^\([a-d]\)', next_text) or
                            # Continuation (short line that doesn't look like new question)
                            (len(next_text) < 50 and 
                             not re.match(r'^\d+[\.\s]', next_text) and 
                             not re.match(r'^[A-H]\.\s+.*', next_text) and
                             not next_text.startswith('Column'))
                        )
                        
                        if is_related:
                            question_content += " " + next_text
                            j += 1
                        else:
                            break
                    
                    # Add the complete question
                    if len(question_content.split()) >= 3:  # Minimum word count
                        questions.append(question_content.strip())
                    
                    # Skip the processed lines
                    i = j
                else:
                    i += 1
            
            # Post-process questions to remove any remaining instructions and categorize
            filtered_questions = []
            question_types = {
                'multiple_choice': [],
                'true_false': [],
                'matching': [], 
                'short_answer': [],
                'long_answer': [],
                'essay': [],
                'other': []
            }
            
            for question in questions:
                # Clean up the question text
                clean_question = question.strip()
                
                # Skip if it's still instruction-like
                if is_instruction_text(clean_question):
                    continue
                
                # Skip if too short to be a real question
                if len(clean_question.split()) < 3:
                    continue
                
                # Skip if it's just formatting or numbers
                if re.match(r'^[\d\s\.\-\(\)]+$', clean_question):
                    continue
                
                filtered_questions.append(clean_question)
                
                # Categorize question type (improved logic)
                question_lower = clean_question.lower()
                
                # Multiple Choice - contains options a), b), c), d) or similar
                if (re.search(r'[a-d]\)|[A-D]\)|^\d+\.\s+.*\b(which|what).*following', clean_question, re.IGNORECASE) or
                    re.search(r'(a\.|b\.|c\.|d\.)\s+\w+', clean_question)):
                    question_types['multiple_choice'].append(clean_question)
                
                # True/False - questions that can be answered with true/false
                elif (re.search(r'^\d+\s+(is|are|does|do|can|will|should|has|have)\s+.*[?.]', clean_question, re.IGNORECASE) or
                      re.search(r'\b(true|false|T/F|correct|incorrect)\b', clean_question, re.IGNORECASE) or
                      # Statements that can be true/false (declarative sentences numbered) - but NOT questions
                      (re.match(r'^\d+\s+[A-Z].*\.$', clean_question) and 
                       len(clean_question.split()) < 12 and
                       not re.search(r'(what|list|describe|explain|how|why|when|where)', clean_question, re.IGNORECASE))):
                    question_types['true_false'].append(clean_question)
                
                # Matching - contains matching keywords or column references
                elif (re.search(r'(match|column A|column B|something \w+)', clean_question, re.IGNORECASE) or
                      re.search(r'^\d+\.\s+(something|cold|hot|sour|sweet)', clean_question, re.IGNORECASE) or
                      # Single letter options (A., B., C., etc.) from matching exercises
                      re.match(r'^[A-H]\.\s+\w+', clean_question, re.IGNORECASE)):
                    question_types['matching'].append(clean_question)
                
                # Essay questions - longer descriptive questions  
                elif re.search(r'(describe|explain|discuss|analyze|compare|contrast)', clean_question, re.IGNORECASE):
                    if len(clean_question.split()) > 12:  # Lowered threshold for essay
                        question_types['essay'].append(clean_question)
                    else:
                        question_types['long_answer'].append(clean_question)
                
                # Short answer questions
                elif (re.search(r'(list|name|identify|define|what is|what are|what)', clean_question, re.IGNORECASE) and
                      len(clean_question.split()) < 15):
                    question_types['short_answer'].append(clean_question)
                
                # Long answer questions
                elif len(clean_question.split()) > 8 and clean_question.endswith('?'):
                    question_types['long_answer'].append(clean_question)
                
                # Default to other
                else:
                    question_types['other'].append(clean_question)
            
            # Extract MCQ options from both standalone paragraphs and within question text
            mcq_options = []
            
            # First, get standalone options
            for para in doc.paragraphs:
                text = para.text.strip()
                if (re.match(r'^[A-D]\)|^[a-d]\)|^\([A-D]\)|^\([a-d]\)', text) and 
                    not is_instruction_text(text) and len(text) > 5):
                    mcq_options.append(text)
            
            # Also extract options from within MCQ question text
            for mcq_question in question_types['multiple_choice']:
                # Find options in a. format (simple and reliable)
                options_dot = re.findall(r'[a-d]\.\s+\w+', mcq_question, re.IGNORECASE)
                mcq_options.extend(options_dot)
                
                # Find options in a) format  
                options_paren = re.findall(r'[a-d]\)\s+\w+', mcq_question, re.IGNORECASE)
                mcq_options.extend(options_paren)
            
            logger.info(f"Extracted {len(filtered_questions)} questions from {os.path.basename(file_path)}")
            logger.info(f"Question types: MCQ={len(question_types['multiple_choice'])}, T/F={len(question_types['true_false'])}, "
                       f"Matching={len(question_types['matching'])}, Short={len(question_types['short_answer'])}, "
                       f"Long={len(question_types['long_answer'])}, Essay={len(question_types['essay'])}, "
                       f"Other={len(question_types['other'])}")
            
            return {
                "questions": filtered_questions,
                "question_types": question_types,
                "mcq_options": mcq_options,
                "total_questions": len(filtered_questions),
                "extraction_time": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error extracting questions from {file_path}: {str(e)}")
            return {
                "questions": [],
                "mcq_options": [],
                "total_questions": 0,
                "error": str(e)
            }
    
    def calculate_tfidf_similarity(self, text1: str, text2: str) -> float:
        """Calculate TF-IDF similarity between two texts"""
        try:
            documents = [text1, text2]
            tfidf_matrix = self.tfidf_vectorizer.fit_transform(documents)
            similarity_matrix = cosine_similarity(tfidf_matrix)
            return float(similarity_matrix[0][1])
        except Exception as e:
            logger.error(f"Error calculating TF-IDF similarity: {str(e)}")
            return 0.0
    
    def calculate_semantic_similarity(self, text1: str, text2: str) -> float:
        """Calculate semantic similarity (placeholder for now)"""
        # For now, return a simulated semantic similarity
        # In production, you would use sentence-transformers or similar
        try:
            # Simple word overlap-based similarity as placeholder
            words1 = set(text1.lower().split())
            words2 = set(text2.lower().split())
            intersection = words1.intersection(words2)
            union = words1.union(words2)
            
            if len(union) == 0:
                return 0.0
            
            jaccard_similarity = len(intersection) / len(union)
            return min(jaccard_similarity * 1.5, 1.0)  # Boost semantic score slightly
            
        except Exception as e:
            logger.error(f"Error calculating semantic similarity: {str(e)}")
            return 0.0
    
    def compare_documents(self, doc1_content: Dict, doc2_content: Dict) -> Dict[str, Any]:
        """Compare two documents using TF-IDF and semantic analysis"""
        questions1 = " ".join(doc1_content.get("questions", []))
        questions2 = " ".join(doc2_content.get("questions", []))
        
        if not questions1 or not questions2:
            return {
                "similarity_score": 0.0,
                "method_used": "none",
                "details": "One or both documents have no extractable questions",
                "processing_time": 0.0
            }
        
        start_time = time.time()
        
        # Step 1: TF-IDF Similarity
        tfidf_score = self.calculate_tfidf_similarity(questions1, questions2)
        
        method_used = "tfidf"
        final_score = tfidf_score
        
        # Step 2: Semantic Analysis (if TF-IDF < threshold)
        if tfidf_score < self.similarity_threshold:
            semantic_score = self.calculate_semantic_similarity(questions1, questions2)
            final_score = max(tfidf_score, semantic_score)
            method_used = "tfidf_semantic"
        
        processing_time = time.time() - start_time
        
        # Find potentially matching questions
        matching_questions = self.find_matching_questions(
            doc1_content.get("questions", []), 
            doc2_content.get("questions", [])
        )
        
        return {
            "similarity_score": round(final_score, 3),
            "tfidf_score": round(tfidf_score, 3),
            "semantic_score": round(self.calculate_semantic_similarity(questions1, questions2), 3) if tfidf_score < self.similarity_threshold else None,
            "method_used": method_used,
            "matching_questions": matching_questions,
            "processing_time": round(processing_time, 3)
        }
    
    def find_matching_questions(self, questions1: List[str], questions2: List[str]) -> List[Dict]:
        """Find potentially matching questions between documents"""
        matches = []
        
        for i, q1 in enumerate(questions1):
            for j, q2 in enumerate(questions2):
                similarity = self.calculate_tfidf_similarity(q1, q2)
                if similarity > 0.6:  # Lower threshold for individual questions
                    matches.append({
                        "question1_index": i,
                        "question2_index": j,
                        "question1": q1[:100] + "..." if len(q1) > 100 else q1,
                        "question2": q2[:100] + "..." if len(q2) > 100 else q2,
                        "similarity": round(similarity, 3)
                    })
        
        # Sort by similarity score (highest first)
        matches.sort(key=lambda x: x["similarity"], reverse=True)
        return matches[:10]  # Return top 10 matches
    
    def analyze_session(self, session_id: str) -> Dict[str, Any]:
        """Analyze all documents in a session and create similarity matrix"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        uploaded_files_path = os.path.join(session_path, "uploaded_files")
        extracted_content_path = os.path.join(session_path, "extracted_content")
        results_path = os.path.join(session_path, "analysis_results")
        
        if not os.path.exists(uploaded_files_path):
            raise ValueError(f"Session {session_id} not found")
        
        # Get all uploaded files
        uploaded_files = [f for f in os.listdir(uploaded_files_path) if f.endswith('.docx')]
        
        if len(uploaded_files) < 2:
            raise ValueError("At least 2 documents required for comparison")
        
        # Extract content from all documents
        extracted_content = {}
        for filename in uploaded_files:
            file_path = os.path.join(uploaded_files_path, filename)
            content = self.extract_questions_from_docx(file_path)
            extracted_content[filename] = content
            
            # Save extracted content
            content_file = os.path.join(extracted_content_path, f"{filename}_questions.json")
            with open(content_file, 'w') as f:
                json.dump(content, f, indent=2)
        
        # Create similarity matrix
        similarity_matrix = []
        detailed_comparisons = {}
        
        # Initialize matrix with filenames
        filenames = list(extracted_content.keys())
        matrix_data = {
            "filenames": filenames,
            "matrix": []
        }
        
        # Calculate similarities
        for i, file1 in enumerate(filenames):
            row = []
            for j, file2 in enumerate(filenames):
                if i == j:
                    # Same file - similarity = 1.0
                    row.append(1.0)
                elif f"{file1}_vs_{file2}" in detailed_comparisons or f"{file2}_vs_{file1}" in detailed_comparisons:
                    # Already calculated (symmetric)
                    existing_key = f"{file2}_vs_{file1}" if f"{file2}_vs_{file1}" in detailed_comparisons else f"{file1}_vs_{file2}"
                    row.append(detailed_comparisons[existing_key]["similarity_score"])
                else:
                    # Calculate new comparison
                    comparison = self.compare_documents(
                        extracted_content[file1], 
                        extracted_content[file2]
                    )
                    detailed_comparisons[f"{file1}_vs_{file2}"] = comparison
                    row.append(comparison["similarity_score"])
            
            matrix_data["matrix"].append(row)
        
        # Save results
        matrix_file = os.path.join(results_path, "similarity_matrix.json")
        with open(matrix_file, 'w') as f:
            json.dump(matrix_data, f, indent=2)
        
        details_file = os.path.join(results_path, "detailed_comparisons.json")
        with open(details_file, 'w') as f:
            json.dump(detailed_comparisons, f, indent=2)
        
        # Update session info
        session_info_path = os.path.join(session_path, "session_info.json")
        with open(session_info_path, 'r') as f:
            session_info = json.load(f)
        
        session_info["analysis_completed"] = True
        session_info["analysis_time"] = datetime.now().isoformat()
        session_info["total_comparisons"] = len(detailed_comparisons)
        
        with open(session_info_path, 'w') as f:
            json.dump(session_info, f, indent=2)
        
        logger.info(f"Analysis completed for session {session_id}")
        
        return {
            "session_id": session_id,
            "similarity_matrix": matrix_data,
            "detailed_comparisons": detailed_comparisons,
            "files_analyzed": len(filenames),
            "total_comparisons": len(detailed_comparisons)
        }
    
    def cleanup_session(self, session_id: str) -> bool:
        """Remove a specific session and all its files"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        
        if os.path.exists(session_path):
            try:
                shutil.rmtree(session_path)
                logger.info(f"Cleaned up session: {session_id}")
                return True
            except Exception as e:
                logger.error(f"Error cleaning up session {session_id}: {str(e)}")
                return False
        
        return False
    
    def cleanup_old_sessions(self, max_age_hours: int = 2) -> int:
        """Remove sessions older than specified hours"""
        if not os.path.exists(self.temp_base_dir):
            return 0
        
        cutoff_time = time.time() - (max_age_hours * 3600)
        cleaned_count = 0
        
        for session_folder in os.listdir(self.temp_base_dir):
            if session_folder.startswith("similarity_"):
                session_path = os.path.join(self.temp_base_dir, session_folder)
                if os.path.isdir(session_path):
                    if os.path.getctime(session_path) < cutoff_time:
                        if self.cleanup_session(session_folder):
                            cleaned_count += 1
        
        logger.info(f"Cleaned up {cleaned_count} old sessions")
        return cleaned_count
    
    def get_session_results(self, session_id: str) -> Dict[str, Any]:
        """Get analysis results for a session"""
        session_path = os.path.join(self.temp_base_dir, session_id)
        results_path = os.path.join(session_path, "analysis_results")
        
        if not os.path.exists(results_path):
            raise ValueError(f"Results not found for session {session_id}")
        
        # Load similarity matrix
        matrix_file = os.path.join(results_path, "similarity_matrix.json")
        with open(matrix_file, 'r') as f:
            similarity_matrix = json.load(f)
        
        # Load detailed comparisons
        details_file = os.path.join(results_path, "detailed_comparisons.json")
        with open(details_file, 'r') as f:
            detailed_comparisons = json.load(f)
        
        # Load session info
        session_info_path = os.path.join(session_path, "session_info.json")
        with open(session_info_path, 'r') as f:
            session_info = json.load(f)
        
        return {
            "session_info": session_info,
            "similarity_matrix": similarity_matrix,
            "detailed_comparisons": detailed_comparisons
        }
